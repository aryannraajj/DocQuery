"""Document processing module for extracting text from various formats."""
import PyPDF2
from pathlib import Path
from typing import Dict, Tuple
import re

class DocumentProcessor:
    """Process documents and extract text content."""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean extracted text by removing extra whitespace and special characters."""
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-\'"]+', '', text)
        return text.strip()
    
    @staticmethod
    def extract_from_pdf(file_path: Path) -> Tuple[str, Dict]:
        """
        Extract text from PDF file.
        
        Returns:
            Tuple of (text_content, metadata)
        """
        text_content = []
        metadata = {"pages": 0, "format": "pdf"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"[Page {page_num}]\n{text}")
                        
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")
        
        full_text = "\n\n".join(text_content)
        return DocumentProcessor.clean_text(full_text), metadata
    
    @staticmethod
    def extract_from_txt(file_path: Path) -> Tuple[str, Dict]:
        """
        Extract text from TXT file.
        
        Returns:
            Tuple of (text_content, metadata)
        """
        metadata = {"format": "txt"}
        
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    metadata["encoding"] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode text file with common encodings")
                
        except Exception as e:
            raise ValueError(f"Error reading TXT: {str(e)}")
        
        return DocumentProcessor.clean_text(text), metadata
    
    @staticmethod
    def extract_from_docx(file_path: Path) -> Tuple[str, Dict]:
        """
        Extract text from DOCX file.
        
        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            metadata = {"format": "docx", "paragraphs": len(doc.paragraphs)}
            
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            full_text = "\n\n".join(text_content)
            return DocumentProcessor.clean_text(full_text), metadata
            
        except ImportError:
            raise ValueError("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")
    
    @classmethod
    def process_document(cls, file_path: Path) -> Tuple[str, Dict]:
        """
        Process a document and extract text based on file extension.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Tuple of (text_content, metadata)
        """
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return cls.extract_from_pdf(file_path)
        elif suffix == '.txt':
            return cls.extract_from_txt(file_path)
        elif suffix == '.docx':
            return cls.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
